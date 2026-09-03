#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
memoir-writer · 导出打印版
将回忆录 Markdown 转为「长辈友好」的打印版 DOCX：
  - 大字号（正文 15pt / 三号感）、宽松行距（1.75）、首行缩进两字
  - 章节分页、黑体标题、宋体/雅黑正文
  - 保留 🥇🥈🥉❓ 证据标记 与 [文学重建] 标注
  - 页脚居中页码；自动标题页

用法：
  python make_print.py --input 回忆录样章.md --output 打印版.docx \
      --title "我的平生履历（回忆录样章）" --author "朱兵 口述/著" \
      --subtitle "温情回忆录 · 征求意见稿"
"""
import argparse
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "✗ 缺少依赖 python-docx。\n"
        "  请先安装：python -m pip install python-docx\n"
        "  或在隔离 venv 中：<venv>/Scripts/python.exe -m pip install python-docx\n"
    )
    sys.exit(2)

# ---- 字体（Windows 常见中文字体）----
BODY_FONT = "微软雅黑"      # 正文：无衬线，长辈清晰
HEAD_FONT = "黑体"          # 标题：醒目
BODY_SIZE = 15              # 正文字号(pt)
TITLE_SIZE = 26
H1_SIZE = 17
H2_SIZE = 15

CJK_MARK = re.compile(r"([🥇🥈🥉❓])")
RECON = "[文学重建]"


def set_cjk(run, font):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)


def add_run(paragraph, text, size=BODY_SIZE, bold=False, italic=False, font=BODY_FONT, color=None):
    """支持 **加粗** 与 emoji 标记的内联渲染。"""
    # 先按 ** 切分
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            r = paragraph.add_run(p[2:-2])
            r.bold = True
            r.font.size = Pt(size)
            set_cjk(r, font)
            if color:
                r.font.color.rgb = color
        else:
            # 逐字符保留 emoji，普通文本正常
            r = paragraph.add_run(p)
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
            set_cjk(r, font)
            if color:
                r.font.color.rgb = color


def style_body(p, size=BODY_SIZE, indent=True, italic=False, align=None, space_after=6):
    p.paragraph_format.line_spacing = 1.75
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 第 ")
    set_cjk(run, BODY_FONT)
    run.font.size = Pt(10)
    # 页码域
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    p._p.append(fld1)
    run2 = p.add_run(" 页 —")
    set_cjk(run2, BODY_FONT)
    run2.font.size = Pt(10)


def add_title_page(doc, title, author, subtitle):
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(t, title, size=TITLE_SIZE, bold=True, font=HEAD_FONT)
    for _ in range(1):
        doc.add_paragraph()
    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(s, subtitle, size=13, font=BODY_FONT, color=RGBColor(0x66, 0x66, 0x66))
    for _ in range(6):
        doc.add_paragraph()
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(a, author, size=13, font=BODY_FONT)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(note, "（本稿由家人协助整理，欢迎提出修改意见）", size=11, font=BODY_FONT,
            color=RGBColor(0x88, 0x88, 0x88))
    doc.add_page_break()


def add_table(doc, rows):
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for c in range(ncol):
            txt = row[c] if c < len(row) else ""
            cell = cells[c]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing = 1.3
            is_head = (i == 0)
            add_run(para, txt, size=11, bold=is_head, font=BODY_FONT)
            if is_head:
                # 表头浅灰底
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'D9E2F3')
                tcPr.append(shd)
    doc.add_paragraph()


def convert(md_text, title, author, subtitle):
    doc = Document()
    # 页面边距
    for sec in doc.sections:
        sec.top_margin = Cm(2.6)
        sec.bottom_margin = Cm(2.6)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(3.0)

    add_title_page(doc, title, author, subtitle)

    lines = md_text.split("\n")
    i = 0
    table_buf = []
    in_quote = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 收集表格
        if stripped.startswith("|"):
            table_buf.append(stripped)
            i += 1
            continue
        else:
            if table_buf:
                # 解析表格（跳过分隔行）
                rows = []
                for tl in table_buf:
                    if re.match(r"^\|[\s:\-|]+\|$", tl):
                        continue
                    cells = [c.strip() for c in tl.strip("|").split("|")]
                    rows.append(cells)
                if rows:
                    add_table(doc, rows)
                table_buf = []

        # 标题
        if stripped.startswith("# ") and not stripped.startswith("##"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, stripped[2:].strip(), size=H1_SIZE + 4, bold=True, font=HEAD_FONT)
            doc.add_paragraph()
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph()
            p = doc.add_paragraph()
            add_run(p, stripped[4:].strip(), size=H2_SIZE, bold=True, font=HEAD_FONT)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_page_break()
            p = doc.add_paragraph()
            add_run(p, stripped[3:].strip(), size=H1_SIZE, bold=True, font=HEAD_FONT)
            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue
        # 引用（诗歌/题记）
        if stripped.startswith(">"):
            in_quote.append(stripped.lstrip("> ").strip())
            i += 1
            continue
        else:
            if in_quote:
                q = doc.add_paragraph()
                q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style_body(q, size=13, indent=False, italic=True, space_after=2)
                for k, ql in enumerate(in_quote):
                    add_run(q, ql, size=13, italic=True, font=BODY_FONT)
                    if k < len(in_quote) - 1:
                        q.add_run("\n")
                in_quote = []
        # 分隔线
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue
        # 列表
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            style_body(p, indent=False)
            add_run(p, re.sub(r"^[-*]\s+", "", stripped), size=BODY_SIZE)
            i += 1
            continue
        # 空行
        if stripped == "":
            i += 1
            continue
        # 普通段落
        p = doc.add_paragraph()
        style_body(p)
        add_run(p, stripped, size=BODY_SIZE)
        i += 1

    # 收尾可能的表格/引用
    if table_buf:
        rows = []
        for tl in table_buf:
            if re.match(r"^\|[\s:\-|]+\|$", tl):
                continue
            rows.append([c.strip() for c in tl.strip("|").split("|")])
        if rows:
            add_table(doc, rows)
    if in_quote:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_body(q, size=13, indent=False, italic=True)
        for k, ql in enumerate(in_quote):
            add_run(q, ql, size=13, italic=True, font=BODY_FONT)
            if k < len(in_quote) - 1:
                q.add_run("\n")

    add_footer(doc)
    return doc


def read_md(path):
    """读取 Markdown，兼容常见中文编码（utf-8 → gbk → latin-1 兜底）。"""
    if not os.path.isfile(path):
        sys.stderr.write("✗ 找不到输入文件：%s\n   请检查路径是否正确（Windows 建议用 C:/Users/... 绝对路径）。\n" % path)
        sys.exit(3)
    for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    sys.stderr.write("✗ 无法以任何常见编码读取：%s\n" % path)
    sys.exit(4)


def main():
    ap = argparse.ArgumentParser(description="回忆录 Markdown → 长辈友好打印版 DOCX")
    ap.add_argument("--input", required=True, help="输入 Markdown 文件路径")
    ap.add_argument("--output", required=True, help="输出 DOCX 文件路径")
    ap.add_argument("--title", default="回忆录", help="标题页主标题")
    ap.add_argument("--author", default="", help="标题页作者署名")
    ap.add_argument("--subtitle", default="", help="标题页副标题")
    args = ap.parse_args()

    if not args.title.strip():
        sys.stderr.write("⚠ 标题为空，已使用默认值「回忆录」。\n")

    md = read_md(args.input)
    if not md.strip():
        sys.stderr.write("✗ 输入文件为空：%s\n" % args.input)
        sys.exit(5)

    try:
        doc = convert(md, args.title, args.author, args.subtitle)
        out_dir = os.path.dirname(os.path.abspath(args.output))
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)
        doc.save(args.output)
    except Exception as e:  # 兜底：不抛 Python 回溯给普通用户
        sys.stderr.write("✗ 生成失败：%s\n   若提示缺少 python-docx，请先 pip install python-docx。\n" % e)
        sys.exit(6)

    print("已生成打印版：", args.output)


if __name__ == "__main__":
    main()
